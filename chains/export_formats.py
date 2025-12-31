#!/usr/bin/env python3
"""
Multi-Format Export Module
Author: Victor Ibhafidon

WHAT IT DOES:
- Exports attack chains to PDF, HTML, and Word/DOCX formats
- Preserves formatting, diagrams, and structure
- Provides professional report layouts
- Supports batch export operations

HOW IT CONNECTS TO THE FRAMEWORK:
- Extends report_generator.py with additional export formats
- Uses chain_analyzer.py for chain data
- Integrates with visualizer.py for diagrams
- Outputs to docs/exports/ folder

USAGE:
    from chains.export_formats import MultiFormatExporter
    
    exporter = MultiFormatExporter()
    exporter.export_to_pdf(chain, "report.pdf")
    exporter.export_to_word(chain, "report.docx")
    exporter.export_to_html(chain, "report.html")
"""

import sys
import os
from typing import Optional
from datetime import datetime

# Add chains directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from chains.chain_analyzer import AttackChain, ImpactLevel
from chains.report_generator import ReportGenerator
from chains.visualizer import generate_markdown_report, generate_mermaid_diagram


class MultiFormatExporter:
    """Exports attack chains to multiple formats"""
    
    def __init__(self):
        """Initialize the exporter"""
        self.report_generator = ReportGenerator()
        self.output_dir = os.path.join(
            os.path.dirname(__file__), '..', 'docs', 'exports'
        )
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_pdf(
        self, 
        chain: AttackChain, 
        output_file: Optional[str] = None
    ) -> str:
        """
        Export chain to PDF format
        
        Args:
            chain: AttackChain to export
            output_file: Output file path (optional)
        
        Returns:
            Path to exported PDF file
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            print("ERROR: reportlab not installed. Install with: pip install reportlab")
            print("Falling back to HTML export...")
            return self.export_to_html(chain, output_file)
        
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in chain.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_file = os.path.join(
                self.output_dir, 
                f"{safe_title}_{timestamp}.pdf"
            )
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Create PDF document
        doc = SimpleDocTemplate(output_file, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#333333',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#555555',
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Title
        story.append(Paragraph(chain.title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        metadata = [
            f"<b>Severity:</b> {chain.severity}",
            f"<b>Impact:</b> {chain.impact.value}",
            f"<b>Report Date:</b> {datetime.now().strftime('%Y-%m-%d')}",
        ]
        if chain.discovered_by:
            metadata.append(f"<b>Discovered By:</b> {chain.discovered_by}")
        if chain.discovered_at:
            metadata.append(f"<b>Discovered At:</b> {chain.discovered_at.strftime('%Y-%m-%d')}")
        
        for meta in metadata:
            story.append(Paragraph(meta, styles['Normal']))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Description
        if chain.description:
            story.append(Paragraph("<b>Description</b>", heading_style))
            story.append(Paragraph(chain.description, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Context
        if chain.context:
            story.append(Paragraph("<b>Context</b>", heading_style))
            story.append(Paragraph(chain.context, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Prerequisites
        if chain.prerequisites:
            story.append(Paragraph("<b>Prerequisites</b>", heading_style))
            for prereq in chain.prerequisites:
                story.append(Paragraph(f"• {prereq}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Attack Chain Steps
        story.append(Paragraph("<b>Attack Chain Steps</b>", heading_style))
        
        for step in chain.steps:
            step_title = f"Step {step.step_number}: {step.vulnerability_type.value}"
            story.append(Paragraph(step_title, styles['Heading3']))
            story.append(Paragraph(f"<b>Description:</b> {step.description}", styles['Normal']))
            
            if step.endpoint:
                story.append(Paragraph(f"<b>Endpoint:</b> <font face='Courier'>{step.endpoint}</font>", styles['Normal']))
            
            if step.payload:
                story.append(Paragraph("<b>Payload:</b>", styles['Normal']))
                story.append(Paragraph(f"<font face='Courier'>{step.payload}</font>", styles['Normal']))
            
            if step.prerequisites:
                story.append(Paragraph("<b>Prerequisites:</b>", styles['Normal']))
                for prereq in step.prerequisites:
                    story.append(Paragraph(f"• {prereq}", styles['Normal']))
            
            if step.outcome:
                story.append(Paragraph(f"<b>Outcome:</b> {step.outcome}", styles['Normal']))
            
            if step.evidence:
                story.append(Paragraph(f"<b>Evidence:</b> {step.evidence}", styles['Normal']))
            
            story.append(Spacer(1, 0.15*inch))
        
        # Recommendations
        recommendations = self.report_generator.generate_recommendations(chain)
        if recommendations:
            story.append(PageBreak())
            story.append(Paragraph("<b>Recommendations</b>", heading_style))
            story.append(Paragraph(recommendations, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        print(f"PDF report exported to: {output_file}")
        return output_file
    
    def export_to_word(
        self, 
        chain: AttackChain, 
        output_file: Optional[str] = None
    ) -> str:
        """
        Export chain to Word/DOCX format
        
        Args:
            chain: AttackChain to export
            output_file: Output file path (optional)
        
        Returns:
            Path to exported DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("ERROR: python-docx not installed. Install with: pip install python-docx")
            print("Falling back to HTML export...")
            return self.export_to_html(chain, output_file)
        
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in chain.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_file = os.path.join(
                self.output_dir, 
                f"{safe_title}_{timestamp}.docx"
            )
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(chain.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Severity: {chain.severity}")
        doc.add_paragraph(f"Impact: {chain.impact.value}")
        doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
        if chain.discovered_by:
            doc.add_paragraph(f"Discovered By: {chain.discovered_by}")
        if chain.discovered_at:
            doc.add_paragraph(f"Discovered At: {chain.discovered_at.strftime('%Y-%m-%d')}")
        
        doc.add_paragraph()  # Blank line
        
        # Description
        if chain.description:
            doc.add_heading('Description', level=1)
            doc.add_paragraph(chain.description)
        
        # Context
        if chain.context:
            doc.add_heading('Context', level=1)
            doc.add_paragraph(chain.context)
        
        # Prerequisites
        if chain.prerequisites:
            doc.add_heading('Prerequisites', level=1)
            for prereq in chain.prerequisites:
                doc.add_paragraph(prereq, style='List Bullet')
        
        # Attack Chain Steps
        doc.add_heading('Attack Chain Steps', level=1)
        
        for step in chain.steps:
            step_title = f"Step {step.step_number}: {step.vulnerability_type.value}"
            doc.add_heading(step_title, level=2)
            
            doc.add_paragraph(f"Description: {step.description}")
            
            if step.endpoint:
                p = doc.add_paragraph(f"Endpoint: ")
                p.add_run(step.endpoint).font.name = 'Courier New'
            
            if step.payload:
                doc.add_paragraph("Payload:")
                p = doc.add_paragraph(step.payload)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
            
            if step.prerequisites:
                doc.add_paragraph("Prerequisites:")
                for prereq in step.prerequisites:
                    doc.add_paragraph(prereq, style='List Bullet')
            
            if step.outcome:
                doc.add_paragraph(f"Outcome: {step.outcome}")
            
            if step.evidence:
                doc.add_paragraph(f"Evidence: {step.evidence}")
            
            doc.add_paragraph()  # Blank line
        
        # Recommendations
        recommendations = self.report_generator.generate_recommendations(chain)
        if recommendations:
            doc.add_page_break()
            doc.add_heading('Recommendations', level=1)
            doc.add_paragraph(recommendations)
        
        # Save document
        doc.save(output_file)
        print(f"Word document exported to: {output_file}")
        return output_file
    
    def export_to_html(
        self, 
        chain: AttackChain, 
        output_file: Optional[str] = None
    ) -> str:
        """
        Export chain to HTML format (enhanced version)
        
        Args:
            chain: AttackChain to export
            output_file: Output file path (optional)
        
        Returns:
            Path to exported HTML file
        """
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in chain.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_file = os.path.join(
                self.output_dir, 
                f"{safe_title}_{timestamp}.html"
            )
        
        # Use existing HTML export from report_generator
        self.report_generator.export_to_html(chain, output_file)
        return output_file
    
    def export_all_formats(
        self, 
        chain: AttackChain, 
        base_name: Optional[str] = None
    ) -> dict:
        """
        Export chain to all available formats
        
        Args:
            chain: AttackChain to export
            base_name: Base name for output files (optional)
        
        Returns:
            Dictionary with format names and file paths
        """
        if not base_name:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in chain.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            base_name = f"{safe_title}_{timestamp}"
        
        results = {}
        
        # Export to HTML (always available)
        html_file = os.path.join(self.output_dir, f"{base_name}.html")
        results['html'] = self.export_to_html(chain, html_file)
        
        # Export to PDF (if available)
        try:
            pdf_file = os.path.join(self.output_dir, f"{base_name}.pdf")
            results['pdf'] = self.export_to_pdf(chain, pdf_file)
        except Exception as e:
            print(f"PDF export failed: {e}")
            results['pdf'] = None
        
        # Export to Word (if available)
        try:
            docx_file = os.path.join(self.output_dir, f"{base_name}.docx")
            results['docx'] = self.export_to_word(chain, docx_file)
        except Exception as e:
            print(f"Word export failed: {e}")
            results['docx'] = None
        
        return results


if __name__ == "__main__":
    # Example usage
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
    
    from chains.chain_analyzer import ChainAnalyzer, ChainStep, VulnerabilityType, ImpactLevel
    
    # Create example chain
    analyzer = ChainAnalyzer()
    chain = analyzer.create_chain(
        title="Example Attack Chain",
        description="An example attack chain for testing multi-format export",
        impact=ImpactLevel.HIGH
    )
    
    step1 = ChainStep(
        1, VulnerabilityType.XSS, 
        "XSS in user profile bio field",
        endpoint="/profile",
        payload="<script>alert('XSS')</script>",
        outcome="XSS payload stored"
    )
    step2 = ChainStep(
        2, VulnerabilityType.SESSION_HIJACKING,
        "Admin views profile, XSS executes in admin context",
        endpoint="/admin/users",
        prerequisites=["XSS payload stored"],
        outcome="Admin session token stolen"
    )
    step3 = ChainStep(
        3, VulnerabilityType.PRIV_ESCALATION,
        "Use stolen admin session token to access admin panel",
        endpoint="/admin/dashboard",
        prerequisites=["Admin session token stolen"],
        outcome="Full admin access"
    )
    
    chain.add_step(step1)
    chain.add_step(step2)
    chain.add_step(step3)
    
    # Export to all formats
    exporter = MultiFormatExporter()
    print("=" * 80)
    print("EXPORTING TO ALL FORMATS")
    print("=" * 80)
    print()
    
    results = exporter.export_all_formats(chain)
    
    print("\nExport Results:")
    for format_name, file_path in results.items():
        if file_path:
            print(f"  {format_name.upper()}: {file_path}")
        else:
            print(f"  {format_name.upper()}: Failed (missing dependencies)")

